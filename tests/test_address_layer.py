"""Тесты адресного слоя: разбор нормативных docx в записи адресов + индекс."""

import docx
import pytest

from src.address_layer import AddressIndex, parse_docx_records


def _make_docx(path):
    d = docx.Document()
    d.add_paragraph("1. Общие положения.")  # пункт
    d.add_paragraph("54. Руководитель организует ТО.")  # пункт
    d.add_paragraph("Таблица 1 - Здания")  # подпись таблицы
    t = d.add_table(rows=3, cols=3)
    t.rows[0].cells[0].text = "Объект защиты"  # шапка — не строка-норма
    t.rows[1].cells[0].text = "1 Здания складов"
    t.rows[1].cells[2].text = "СПС независимо от площади"
    t.rows[2].cells[0].text = "18. Отдельно стоящие до 100 м2"
    t.rows[2].cells[2].text = "Независимо от площади"
    d.save(path)


@pytest.mark.unit
def test_parse_points_and_tables(tmp_path):
    p = tmp_path / "mini.docx"
    _make_docx(str(p))
    recs = parse_docx_records(str(p), doc_id="mini.docx")
    points = {r["num"]: r for r in recs if r["kind"] == "point"}
    assert "54" in points and "ТО" in points["54"]["text"]
    tbl = {(r["table"], r["row"]): r for r in recs if r["kind"] == "table"}
    assert ("1", "18") in tbl
    assert "100 м2" in tbl[("1", "18")]["text"]
    # шапка «Объект защиты» без номера — не строка-норма
    assert ("1", "0") not in tbl


@pytest.mark.unit
def test_unlabeled_table_skipped(tmp_path):
    d = docx.Document()
    t = d.add_table(rows=3, cols=2)  # без подписи «Таблица N»
    t.rows[1].cells[0].text = "1 строка"
    p = tmp_path / "u.docx"
    d.save(str(p))
    recs = parse_docx_records(str(p), doc_id="u.docx")
    assert [r for r in recs if r["kind"] == "table"] == []


@pytest.mark.unit
def test_build_and_lookup():
    recs = [
        {"doc": "d.docx", "kind": "point", "num": "54", "text": "54. текст"},
        {"doc": "d.docx", "kind": "table", "table": "1", "row": "18", "text": "стр 18"},
        {"doc": "d.docx", "kind": "table", "table": "1", "row": "19", "text": "стр 19"},
    ]
    idx = AddressIndex(recs)
    assert idx.lookup("d.docx", "point", "54", None)[0]["text"] == "54. текст"
    assert idx.lookup("d.docx", "table", "1", "18")[0]["text"] == "стр 18"
    assert len(idx.lookup("d.docx", "table", "1", None)) == 2  # вся таблица
    assert idx.lookup("d.docx", "point", "999", None) == []


@pytest.mark.unit
def test_index_json_roundtrip(tmp_path):
    recs = [{"doc": "d.docx", "kind": "point", "num": "1", "text": "1. x"}]
    p = tmp_path / "idx.json"
    AddressIndex(recs).save(p)
    idx = AddressIndex.load(p)
    assert idx.lookup("d.docx", "point", "1", None)[0]["text"] == "1. x"


@pytest.mark.unit
def test_load_missing_file_is_empty(tmp_path):
    idx = AddressIndex.load(tmp_path / "nope.json")
    assert idx.lookup("d.docx", "point", "1", None) == []
